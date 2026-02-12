
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT

# File paths
md_file = "d:/agentic-ai/resume/VaranSinghRohila_Resume_Updated.md"
docx_file = "d:/agentic-ai/resume/VaranSinghRohila_Resume.docx"
pdf_file = "d:/agentic-ai/resume/VaranSinghRohila_Resume.pdf"

def read_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.readlines()

def create_docx(lines, output_path):
    doc = Document()
    
    # Adjust margins to fit one page better
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            # Name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = 'Calibri'
        elif "Jersey City, NJ" in line:
            # Contact Info
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)
        elif line.startswith("## "):
            # Section Headers
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(12)
            run.underline = True
        elif line.startswith("**") and "** (" in line:
             # Project Title with (Category)
             p = doc.add_paragraph()
             text = line.replace("**", "")
             run = p.add_run(text)
             run.bold = True
        elif line.startswith("**") and "|" in line:
             # Job Title | Company | Date
             p = doc.add_paragraph()
             parts = line.split("|")
             # Bold the title and company
             title_company = parts[0].strip().replace("**", "") + " | " + parts[1].strip()
             date = parts[2].strip().replace("*", "")
             
             run = p.add_run(title_company)
             run.bold = True
             p.add_run(" | " + date)
        
        elif line.startswith("**"):
             # Education / other bold lines
             p = doc.add_paragraph()
             text = line.replace("**", "")
             run = p.add_run(text)
             run.bold = True

        elif line.startswith("* "):
            # Bullet points
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            
            # Handle bold text within bullets
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Word document saved to {output_path}")

def create_pdf(lines, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                            rightMargin=50, leftMargin=50,
                            topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'Title', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=24, alignment=1, spaceAfter=6
    )
    contact_style = ParagraphStyle(
        'Contact', parent=styles['Normal'], fontName='Helvetica', fontSize=10, alignment=1, spaceAfter=12
    )
    heading_style = ParagraphStyle(
        'Heading', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12, spaceAfter=6, spaceBefore=12
    )
    # Draw a line under heading? handled by just bolding for now to be simple
    
    subheading_style = ParagraphStyle(
        'SubHeading', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10, spaceAfter=2
    )
    
    bullet_style = ParagraphStyle(
        'Bullet', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leftIndent=12, firstLineIndent=0, spaceAfter=2
    )

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Escape XML characters for ReportLab
        line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        # Replace markdown links with text for PDF
        line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
        
        # Replace **bold** with <b>bold</b> for reportlab
        line_html = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
        # Replace *italic* with <i>italic</i>, carefully avoiding bullet points
        # only match *text* if it's not at the start of the line (bullet) or followed by space
        # A simple way for this specific resume content:
        line_html = re.sub(r'(?<!^)\*(.*?)\*', r'<i>\1</i>', line_html)

        if line.startswith("# "):
            story.append(Paragraph(line[2:], title_style))
        elif "Jersey City, NJ" in line:
            story.append(Paragraph(line_html, contact_style))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], heading_style))
        elif line.startswith("**") and "|" in line:
            # Job Title
            story.append(Paragraph(line_html, subheading_style))
        elif line.startswith("**"):
             story.append(Paragraph(line_html, subheading_style))
        elif line.startswith("* "):
            story.append(Paragraph(line_html[2:], bullet_style))
        else:
            story.append(Paragraph(line_html, styles['Normal']))

    doc.build(story)
    print(f"PDF document saved to {output_path}")

if __name__ == "__main__":
    lines = read_markdown(md_file)
    create_docx(lines, docx_file)
    create_pdf(lines, pdf_file)
